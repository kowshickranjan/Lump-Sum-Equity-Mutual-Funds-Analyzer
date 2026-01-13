# app.py

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import subprocess
import sys
from io import BytesIO

# ----------------------------
# Fund Name Mapping
# ----------------------------
FUND_NAME_MAP = {
    "F001": "ABSL Large Cap", "F002": "Axis Large Cap", "F003": "DSP Large Cap",
    "F004": "HDFC Large Cap", "F005": "HSBC Large Cap", "F006": "ICICI Large Cap",
    "F007": "Kotak Large Cap", "F008": "UTI Large Cap", "F009": "ABSL Mid Cap",
    "F010": "Axis Mid Cap", "F011": "DSP Mid Cap", "F012": "HDFC Mid Cap",
    "F013": "HSBC Mid Cap", "F014": "Kotak Mid Cap", "F015": "UTI Mid Cap",
    "F016": "ABSL Small Cap", "F017": "Axis Small Cap", "F018": "DSP Small Cap",
    "F019": "HDFC Small Cap", "F020": "HSBC Small Cap", "F021": "ICICI Small Cap",
    "F022": "Kotak Small Cap", "F023": "ABSL Lg+Mid", "F024": "Axis Lg+Mid",
    "F025": "DSP Lg+Mid", "F026": "HDFC Lg+Mid", "F027": "HSBC Lg+Mid",
    "F028": "ICICI Lg+Mid", "F029": "Kotak Lg+Mid", "F030": "UTI Lg+Mid",
    "F031": "ABSL Flexi Cap", "F032": "Axis Flexi Cap", "F033": "DSP Flexi Cap",
    "F034": "HDFC Flexi Cap", "F035": "HSBC Flexi Cap", "F036": "Kotak Flexi Cap",
    "F037": "UTI Flexi Cap"
}

def get_fund_display_name(fund_code):
    return FUND_NAME_MAP.get(fund_code, fund_code)

# ----------------------------
# Prophet Install Guard
# ----------------------------
def ensure_prophet_installed():
    try:
        import prophet
    except ImportError:
        st.warning("Prophet not found. Installing...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "prophet"])
            st.success("Prophet installed! Restarting app...")
            st.rerun()
        except Exception as e:
            st.error(f"Failed to install Prophet: {e}")
            st.stop()

ensure_prophet_installed()
from prophet import Prophet

# ----------------------------
# Helper: Plot to BytesIO
# ----------------------------
def plot_to_bytes(fig, dpi=150):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# ----------------------------
# App Setup
# ----------------------------
st.set_page_config(layout="wide", page_title="Lump Sum Equity Mutual Funds Analyzer")
st.title(" Lump Sum Equity Mutual Funds Analyzer ")

uploaded_file = st.file_uploader("Upload your NAV_DATA.csv", type="csv")

if uploaded_file is not None:

    @st.cache_data(show_spinner="Processing data...")
    def preprocess_data(file):
        df = pd.read_csv(file)
        if 'Nav_Date' not in df.columns:
            st.error("Missing 'Nav_Date' column.")
            st.stop()
        df['Nav_Date'] = pd.to_datetime(df['Nav_Date'], format='%d-%m-%Y', errors='coerce')
        df.dropna(subset=['Nav_Date'], inplace=True)
        if df.empty:
            st.error("No valid dates found. Use DD-MM-YYYY.")
            st.stop()
        df.sort_values('Nav_Date', inplace=True)
        df.reset_index(drop=True, inplace=True)

        fund_cols = [col for col in df.columns if col.startswith('F')]
        if not fund_cols:
            st.error("No fund columns (e.g., F001) found.")
            st.stop()

        valid_funds = []
        for col in fund_cols:
            df[col] = pd.to_numeric(df[col], errors='coerce')
            if df[col].notna().any():
                valid_funds.append(col)
            else:
                st.warning(f"Skipping empty fund: {col}")

        if not valid_funds:
            st.error("No valid fund data remains.")
            st.stop()

        prophet_datasets = {}
        for fund in valid_funds:
            fund_df = df[['Nav_Date', fund]].dropna().copy()
            fund_df.rename(columns={fund: 'nav'}, inplace=True)
            fund_df['daily_log_ret'] = np.log(fund_df['nav'] / fund_df['nav'].shift(1))
            fund_df['y'] = fund_df['daily_log_ret'].rolling(window=252, min_periods=100).sum()
            fund_df = fund_df.dropna(subset=['y'])[['Nav_Date', 'y']].rename(columns={'Nav_Date': 'ds'})
            if len(fund_df) < 2:
                continue
            fund_df.set_index('ds', inplace=True)
            fund_df = fund_df.resample('ME').last().reset_index()
            if len(fund_df) >= 2:
                prophet_datasets[fund] = fund_df
        if not prophet_datasets:
            st.error("No funds have enough monthly data.")
            st.stop()
        return prophet_datasets, valid_funds

    try:
        fund_data, all_valid_funds = preprocess_data(uploaded_file)
        st.success(f"✅ Processed {len(fund_data)} funds.")
    except Exception as e:
        st.error(f"Data error: {e}")
        st.stop()

    first_fund = next(iter(fund_data))
    st.subheader("Sample Processed Data")
    st.dataframe(fund_data[first_fund].head())

    # ----------------------------
    # Fund Selection: One by One
    # ----------------------------
    st.header("Select Funds for Analysis (Choose at least 3)")
    selected_funds = set()

    # Allow up to 10 selections or all available
    max_select = min(10, len(all_valid_funds))
    for i in range(max_select):
        key = f"fund_select_{i}"
        options = [f for f in all_valid_funds if f not in selected_funds]
        if not options:
            break
        fund = st.selectbox(
            f"Fund {i+1}",
            options=[""] + options,
            format_func=lambda x: get_fund_display_name(x) if x else "Select a fund",
            key=key
        )
        if fund:
            selected_funds.add(fund)
        else:
            break

    selected_funds = list(selected_funds)
    if len(selected_funds) < 3:
        st.warning("⚠️ Please select at least 3 funds to proceed.")
    else:
        if st.button("🚀 Generate Analysis"):
            with st.spinner("Training models..."):

                @st.cache_data
                def run_forecasts(fund_dict):
                    forecasts = {}
                    all_preds = []
                    for fund, data in fund_dict.items():
                        try:
                            model = Prophet(yearly_seasonality=True, weekly_seasonality=False, daily_seasonality=False, changepoint_prior_scale=0.05)
                            model.fit(data)
                            future = model.make_future_dataframe(periods=12, freq='ME')
                            pred = model.predict(future)
                            forecasts[fund] = pred
                            pred_sub = pred[['ds', 'yhat', 'yhat_lower', 'yhat_upper']].copy()
                            pred_sub['fund'] = fund
                            all_preds.append(pred_sub)
                        except Exception as e:
                            st.error(f"Error in {fund}: {e}")
                    combined = pd.concat(all_preds, ignore_index=True) if all_preds else pd.DataFrame()
                    return forecasts, combined

                subset = {f: fund_data[f] for f in selected_funds}
                forecasts, all_forecasts_df = run_forecasts(subset)

                if not forecasts:
                    st.error("No forecasts generated.")
                else:
                    st.success("✅ Analysis complete!")

                    # --- Summary Table ---
                    st.header("📊 Forecast Summary (Next 12 Months)")
                    summary_rows = []
                    for fund in selected_funds:
                        hist_end = fund_data[fund]['ds'].max()
                        future = forecasts[fund][forecasts[fund]['ds'] > hist_end]
                        if future.empty:
                            continue
                        mean_ret = future['yhat'].mean()
                        best_case = future['yhat_upper'].max()
                        worst_case = future['yhat_lower'].min()
                        volatility = future['yhat'].std()
                        summary_rows.append({
                            "Fund Code": fund,
                            "Fund Name": get_fund_display_name(fund),
                            "Mean Forecasted Return": mean_ret,
                            "Best Case (Max Upper)": best_case,
                            "Worst Case (Min Lower)": worst_case,
                            "Volatility (σ)": volatility
                        })

                    summary_df = pd.DataFrame(summary_rows)
                    if not summary_df.empty:
                        summary_df = summary_df.sort_values("Mean Forecasted Return", ascending=False)
                        summary_df.set_index("Fund Name", inplace=True)
                        st.dataframe(summary_df.style.format({
                            "Mean Forecasted Return": "{:.2%}",
                            "Best Case (Max Upper)": "{:.2%}",
                            "Worst Case (Min Lower)": "{:.2%}",
                            "Volatility (σ)": "{:.2%}"
                        }))
                    else:
                        st.info("No forecast data for summary.")

                    # --- Store Plots with Data Points ---
                    plot_images = []
                    st.header("📈 Individual Fund Forecasts")
                    for fund in selected_funds:
                        if fund in forecasts:
                            hist = fund_data[fund]
                            pred = all_forecasts_df[all_forecasts_df['fund'] == fund]
                            fig, ax = plt.subplots(figsize=(6, 3))
                            # Historical points
                            ax.plot(hist['ds'], hist['y'], color="steelblue", alpha=0.7)
                            ax.scatter(hist['ds'], hist['y'], color="steelblue", s=10, label="Historical")
                            # Forecast line & CI
                            ax.plot(pred['ds'], pred['yhat'], color="crimson", linestyle="--", label="Forecast")
                            ax.fill_between(pred['ds'], pred['yhat_lower'], pred['yhat_upper'], color="crimson", alpha=0.2)
                            # Forecast points (future only)
                            future_pred = pred[pred['ds'] > hist['ds'].max()]
                            ax.scatter(future_pred['ds'], future_pred['yhat'], color="crimson", s=10)
                            ax.set_title(f"{get_fund_display_name(fund)}", fontsize=10)
                            ax.set_ylabel("Rolling 1-Year Log Return", fontsize=8)
                            ax.legend(fontsize=7)
                            ax.grid(True, linestyle="--", alpha=0.5)
                            plt.xticks(rotation=30, fontsize=7)
                            plt.yticks(fontsize=7)
                            plt.tight_layout()
                            img_bytes = plot_to_bytes(fig, dpi=150)
                            plot_images.append((get_fund_display_name(fund), img_bytes))
                            st.pyplot(fig)

                    # --- Allocation Section ---
                    st.header("💼 Lump-Sum Investment Allocation Planner")
                    investment_lakhs = st.number_input(
                        "Enter your lump-sum investment amount (₹ in Lakhs):",
                        min_value=1.0,
                        max_value=10.0,
                        value=5.0,
                        step=0.5,
                        format="%.1f"
                    )
                    total_investment = investment_lakhs * 100000

                    fund_metrics = []
                    for fund in selected_funds:
                        hist_end = fund_data[fund]['ds'].max()
                        future = forecasts[fund][forecasts[fund]['ds'] > hist_end]
                        if future.empty:
                            continue
                        mean_ret = future['yhat'].mean()
                        vol = future['yhat'].std()
                        fund_metrics.append({
                            'fund': fund,
                            'mean_return': mean_ret,
                            'volatility': vol
                        })

                    alloc_summary_df = pd.DataFrame()
                    returns_comparison_df = pd.DataFrame()
                    if fund_metrics:
                        metrics_df = pd.DataFrame(fund_metrics).set_index('fund')
                        n = len(metrics_df)
                        equal_weights = pd.Series([1/n] * n, index=metrics_df.index)
                        equal_return = (equal_weights * metrics_df['mean_return']).sum()

                        returns = metrics_df['mean_return']
                        min_ret = returns.min()
                        shifted = returns - min_ret + 1e-6 if min_ret < 0 else returns + 1e-6
                        high_ret_weights = shifted / shifted.sum()
                        high_ret_return = (high_ret_weights * returns).sum()

                        vols = metrics_df['volatility'].copy()
                        vols = vols.replace(0, vols[vols > 0].min() if vols[vols > 0].any() else 1e-6)
                        inv_vol = 1 / vols
                        low_risk_weights = inv_vol / inv_vol.sum()
                        low_risk_return = (low_risk_weights * returns).sum()

                        strategies = {
                            "Equal Weight": {"weights": equal_weights, "return": equal_return},
                            "High Return Focus": {"weights": high_ret_weights, "return": high_ret_return},
                            "Low Risk Focus": {"weights": low_risk_weights, "return": low_risk_return}
                        }

                        all_alloc_data = []
                        comp_data = []
                        for name, strat in strategies.items():
                            alloc_rupees = (strat["weights"] * total_investment).round(0)
                            weights_pct = (strat["weights"] * 100).round(2)
                            fund_returns = metrics_df.loc[alloc_rupees.index, 'mean_return']
                            for fund_code in alloc_rupees.index:
                                all_alloc_data.append({
                                    "Strategy": name,
                                    "Fund Name": get_fund_display_name(fund_code),
                                    "Allocation (₹)": alloc_rupees[fund_code],
                                    "Weight (%)": weights_pct[fund_code],
                                    "Fund Return": fund_returns[fund_code]
                                })
                            comp_data.append({"Strategy": name, "Portfolio Return": strat["return"]})

                        alloc_summary_df = pd.DataFrame(all_alloc_data)
                        returns_comparison_df = pd.DataFrame(comp_data)

                        for name in strategies.keys():
                            st.subheader(f"🔹 {name} Strategy")
                            strat_df = alloc_summary_df[alloc_summary_df["Strategy"] == name].copy()
                            portfolio_return = returns_comparison_df[returns_comparison_df["Strategy"] == name]["Portfolio Return"].iloc[0]
                            st.write(f"**Expected Portfolio Return**: {portfolio_return:.2%}")
                            st.dataframe(strat_df[["Fund Name", "Allocation (₹)", "Weight (%)", "Fund Return"]].style.format({
                                "Allocation (₹)": "₹{:,.0f}",
                                "Weight (%)": "{:.2f}%",
                                "Fund Return": "{:.2%}"
                            }), use_container_width=True)

                        # --- Comparison Bar Chart ---
                        st.header("📈 Strategy Return Comparison")
                        fig, ax = plt.subplots(figsize=(6, 3))
                        bars = ax.bar(returns_comparison_df["Strategy"], returns_comparison_df["Portfolio Return"],
                                      color=['#4E79A7', '#F28E2B', '#E15759'])
                        ax.set_ylabel("Expected Portfolio Return", fontsize=9)
                        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda y, _: '{:.1%}'.format(y)))
                        for bar, val in zip(bars, returns_comparison_df["Portfolio Return"]):
                            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.001,
                                    f'{val:.2%}', ha='center', va='bottom', fontweight='bold', fontsize=8)
                        plt.xticks(rotation=0, fontsize=8)
                        plt.yticks(fontsize=8)
                        plt.tight_layout()
                        comp_plot_bytes = plot_to_bytes(fig, dpi=150)
                        st.pyplot(fig)

                    # ---------------------------------------------------------
                    # 📥 DOWNLOAD FULL REPORT AS WORD (.docx)
                    # ---------------------------------------------------------
                    st.header("📥 Download Full Report as Word Document")

                    try:
                        from docx import Document
                        from docx.shared import Inches, Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        doc = Document()
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Arial'
                        font.size = Pt(10)

                        # Title
                        title = doc.add_heading('Lump Sum Equity Mutual Funds Analyzer', 0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph(f"Investment Amount: ₹{investment_lakhs:,.1f} Lakhs")
                        doc.add_paragraph()

                        # --- 1. Summary Table ---
                        if not summary_df.empty:
                            doc.add_heading("1. Forecast Summary", level=1)
                            table = doc.add_table(rows=1, cols=5)
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            headers = ["Fund Name", "Mean Return", "Best Case", "Worst Case", "Volatility"]
                            for i, header in enumerate(headers):
                                hdr_cells[i].text = header
                                hdr_cells[i].paragraphs[0].runs[0].font.bold = True

                            for _, row in summary_df.reset_index().iterrows():
                                cells = table.add_row().cells
                                cells[0].text = str(row["Fund Name"])
                                cells[1].text = f"{row['Mean Forecasted Return']:.2%}"
                                cells[2].text = f"{row['Best Case (Max Upper)']:.2%}"
                                cells[3].text = f"{row['Worst Case (Min Lower)']:.2%}"
                                cells[4].text = f"{row['Volatility (σ)']:.2%}"
                            doc.add_paragraph()

                        # --- 2. Forecast Plots ---
                        if plot_images:
                            doc.add_heading("2. Forecast Plots", level=1)
                            for fund_name, img_bytes in plot_images:
                                doc.add_paragraph(f"{fund_name}", style='Heading3')
                                doc.add_picture(img_bytes, width=Inches(6))
                                doc.add_paragraph()

                        # --- 3. Allocation Tables ---
                        if not alloc_summary_df.empty:
                            doc.add_heading("3. Allocation Strategies", level=1)
                            for strategy in alloc_summary_df["Strategy"].unique():
                                strat_df = alloc_summary_df[alloc_summary_df["Strategy"] == strategy]
                                port_ret = returns_comparison_df[returns_comparison_df["Strategy"] == strategy]["Portfolio Return"].iloc[0]
                                doc.add_paragraph(f"{strategy} (Return: {port_ret:.2%})", style='Heading2')
                                alloc_table = doc.add_table(rows=1, cols=4)
                                alloc_table.style = 'Table Grid'
                                hdr = alloc_table.rows[0].cells
                                for i, h in enumerate(["Fund Name", "Allocation (₹)", "Weight (%)", "Fund Return"]):
                                    hdr[i].text = h
                                    hdr[i].paragraphs[0].runs[0].font.bold = True

                                for _, r in strat_df.iterrows():
                                    c = alloc_table.add_row().cells
                                    c[0].text = str(r["Fund Name"])
                                    c[1].text = f"₹{r['Allocation (₹)']:,.0f}"
                                    c[2].text = f"{r['Weight (%)']:.2f}%"
                                    c[3].text = f"{r['Fund Return']:.2%}"
                                doc.add_paragraph()

                        # --- 4. Comparison Plot ---
                        if 'comp_plot_bytes' in locals():
                            doc.add_heading("4. Strategy Comparison", level=1)
                            doc.add_picture(comp_plot_bytes, width=Inches(6))

                        # Save and offer download
                        word_buffer = BytesIO()
                        doc.save(word_buffer)
                        word_buffer.seek(0)

                        st.download_button(
                            label="📘 Download Full Report as Word (.docx)",
                            data=word_buffer.read(),
                            file_name=f"LumpSum_Equity_Analysis.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                    except ImportError:
                        st.warning("Install python-docx: pip install python-docx")
                        csv_data = alloc_summary_df.to_csv(index=False).encode('utf-8')
                        st.download_button(
                            label="📄 Download Allocation as CSV (Fallback)",
                            data=csv_data,
                            file_name="allocation_fallback.csv",
                            mime="text/csv"
                        )

else:
    st.info("👆 Upload your `NAV_DATA.csv` to begin.")